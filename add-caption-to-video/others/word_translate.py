from docx import Document
import re
from docx.shared import Inches
import boto3


language_list = ['en', 'es', 'fr']
file_name = '/Users/mac/Desktop/test.docx'

def split_text(_translate, str_text , language):
    """
    将字符串截断以后 进行语言转换
    :param _translate:
    :param str_text:
    :return:
    """
    str_items = str_text.split('：')
    if len(str_items) > 1 and str_items[0] != '人物':
        result = _translate.translate_text(Text=str_items[1], SourceLanguageCode="zh", TargetLanguageCode=language)
        return result['TranslatedText'].strip('\n')
    else:
        return None


def translate_doc(document_old, new_file_name, language):
    """
    按照指定语言， 进行翻译
    :param document_old:
    :param document_new:
    :param language:
    :return:
    """

    translate = boto3.client(service_name='translate', region_name='us-east-1', use_ssl=True)

    for paragraph in document_old.paragraphs:
        result_text = split_text(translate, paragraph.text, language)  # 打印各段落内容文本
        print(paragraph.text)
        document_new.add_paragraph(paragraph.text, style=paragraph.style)
        if result_text != None:
            print(result_text)
            document_new.add_paragraph(result_text, style=paragraph.style)

    document_new.save(new_file_name)


if __name__ == "__main__":


    # 需要翻译的文件
    document_old = Document(file_name)  # 打开文件demo.docx

    for _language in language_list:
        # 生成新文件的名称
        item_list = file_name.rsplit('.', 1)
        new_file_name = item_list[0] + '_' + _language + '.' + item_list[1]
        print(new_file_name)

        # 新生成的文件
        document_new = Document()

        translate_doc(document_old, new_file_name, _language)
